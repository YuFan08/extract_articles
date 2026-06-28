from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from extract_articles import extract, iter_articles


def test_split_and_export_two_articles():
    with TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        source = tmp / "source.docx"
        doc = Document()
        doc.add_paragraph("One", style="Heading 1")
        doc.add_paragraph("English text.")
        doc.add_paragraph("中文正文。")
        doc.add_paragraph("Two", style="Heading 1")
        doc.add_paragraph("More English.")
        doc.save(source)

        assert [title for title, _ in iter_articles(Document(source))] == ["One", "Two"]
        assert extract(source, tmp / "out", make_pdf=False) == 2
        assert len(list((tmp / "out" / "word").glob("*.docx"))) == 2




def test_output_dir_cannot_be_existing_file():
    with TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        source = tmp / "source.docx"
        doc = Document()
        doc.add_paragraph("One", style="Heading 1")
        doc.add_paragraph("English text.")
        doc.save(source)

        try:
            extract(source, source, make_pdf=False)
        except ValueError as exc:
            assert "输出路径不能是已有文件" in str(exc)
        else:
            raise AssertionError("expected ValueError")



def test_safe_stem_replaces_smart_quotes():
    from extract_articles import safe_stem

    assert safe_stem("China’s currency: Tricky troika", 2) == "002-China_s currency_ Tricky troika"
