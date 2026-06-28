from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CJK_RE = re.compile(r"[\u3400-\u9fff]")
BAD_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]+')


def progress(current: int, total: int, message: str) -> None:
    width = 24
    done = int(width * current / total) if total else width
    bar = "#" * done + "-" * (width - done)
    print(f"[{bar}] {current}/{total} {message}", flush=True)


def safe_stem(text: str, index: int) -> str:
    name = BAD_FILENAME_CHARS.sub("_", text).strip().strip(".")
    return f"{index:03d}-{name[:80] or 'article'}"


def set_run_font(run, latin: str, east_asia: str, size_pt: int, color: str = "000000") -> None:
    run.font.name = latin
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def shade(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_body_paragraph_format(paragraph, is_zh: bool) -> None:
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1 if is_zh else 1.25
    paragraph.paragraph_format.first_line_indent = Pt(28)
    if not is_zh:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def iter_articles(doc: Document):
    title = None
    body = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue
        if para.style.name.startswith("Heading 1"):
            if title and body:
                yield title, body
            title, body = text, []
        elif title:
            body.append(text)
    if title and body:
        yield title, body


def write_article(title: str, body: list[str], out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(14)
    title_run = title_p.add_run(title)
    title_run.bold = True
    if CJK_RE.search(title):
        set_run_font(title_run, "Microsoft YaHei", "微软雅黑", 18, "1F2937")
    else:
        set_run_font(title_run, "Cooper Black", "Cooper Black", 18, "1F2937")

    for text in body:
        is_zh = bool(CJK_RE.search(text))
        p = doc.add_paragraph()
        set_body_paragraph_format(p, is_zh)
        shade(p, "DDEBFF" if is_zh else "E2F0D9")
        run = p.add_run(text)
        if is_zh:
            set_run_font(run, "KaiTi", "楷体", 14, "0F172A")
        else:
            set_run_font(run, "Times New Roman", "Times New Roman", 14, "111827")

    doc.save(out_path)


def find_soffice() -> str | None:
    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    for path in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if Path(path).exists():
            return path
    return None


def convert_with_soffice(docx_path: Path, pdf_dir: Path) -> bool:
    soffice = find_soffice()
    if not soffice:
        return False
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(docx_path)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return (pdf_dir / f"{docx_path.stem}.pdf").exists()


def ps_quote(value: Path) -> str:
    return "'" + str(value).replace("'", "''") + "'"


def convert_with_word(docx_path: Path, pdf_path: Path) -> bool:
    if sys.platform != "win32":
        return False
    docx = ps_quote(docx_path)
    pdf = ps_quote(pdf_path)
    ps = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
    $doc = $word.Documents.Open({docx}, $false, $true)
    $doc.ExportAsFixedFormat({pdf}, 17)
    $doc.Close([ref]$false)
}} finally {{
    $word.Quit()
}}
"""
    proc = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return proc.returncode == 0 and pdf_path.exists()


def convert_to_pdf(docx_path: Path, pdf_dir: Path) -> None:
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    if convert_with_soffice(docx_path, pdf_dir) or convert_with_word(docx_path.resolve(), pdf_path.resolve()):
        return
    raise RuntimeError(f"PDF 转换失败：{docx_path.name}。请确认已安装 LibreOffice 或 Microsoft Word。")


def validate_paths(input_docx: Path, output_dir: Path) -> None:
    if not input_docx.exists():
        raise ValueError(f"输入文件不存在：{input_docx}")
    if not input_docx.is_file():
        raise ValueError(f"输入路径不是文件：{input_docx}")
    if output_dir.exists() and not output_dir.is_dir():
        suggestion = output_dir.with_suffix("")
        raise ValueError(f"输出路径不能是已有文件：{output_dir}\n请改用文件夹，例如：-o \"{suggestion}\"")


def extract(input_docx: Path, output_dir: Path, max_articles: int | None = None, make_pdf: bool = True) -> int:
    validate_paths(input_docx, output_dir)
    word_dir = output_dir / "word"
    pdf_dir = output_dir / "pdf"

    source = Document(input_docx)
    articles = list(iter_articles(source))
    if max_articles:
        articles = articles[:max_articles]
    if not articles:
        raise ValueError("没有找到文章标题。请确认文档使用 Heading 1 作为每篇文章标题。")

    word_dir.mkdir(parents=True, exist_ok=True)
    if make_pdf:
        pdf_dir.mkdir(parents=True, exist_ok=True)

    total = len(articles)
    for count, (title, body) in enumerate(articles, 1):
        stem = safe_stem(title, count)
        docx_path = word_dir / f"{stem}.docx"
        progress(count, total, f"正在提取：{title}")
        write_article(title, body, docx_path)
        if make_pdf:
            progress(count, total, f"正在转换 PDF：{title}")
            convert_to_pdf(docx_path, pdf_dir)
    return total


def main() -> None:
    parser = argparse.ArgumentParser(description="Split a bilingual Economist DOCX into styled DOCX and PDF files.")
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("-o", "--output-dir", type=Path)
    parser.add_argument("--max-articles", type=int, help="Only export the first N articles, useful for checking format.")
    parser.add_argument("--no-pdf", action="store_true", help="Only export Word files.")
    args = parser.parse_args()

    output_dir = args.output_dir or args.input_docx.with_name(f"{args.input_docx.stem}_拆分文章")
    try:
        count = extract(args.input_docx, output_dir, args.max_articles, make_pdf=not args.no_pdf)
    except (OSError, ValueError, RuntimeError) as exc:
        print(f"错误：{exc}", file=sys.stderr)
        raise SystemExit(2)
    kinds = "Word 和 PDF" if not args.no_pdf else "Word"
    print(f"完成：已导出 {count} 篇文章为 {kinds}，位置：{output_dir}")


if __name__ == "__main__":
    main()
